# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import json
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"E:\class_work\TrafficSign_CNN_GTSRB")
OLD_TEMPLATE = Path(r"C:\Users\Lenovo\Desktop\xeizuo\机器学习\大作业.docx")
OUT = ROOT / "docs" / "基于ResNet18-FPN的交通标志识别系统大作业.docx"
FIG_DIR = ROOT / "outputs" / "figures"
LOG_PATH = ROOT / "outputs" / "logs" / "training_log.csv"
VAL_REPORT = ROOT / "outputs" / "reports" / "validation_classification_report.txt"
TEST_REPORT = ROOT / "outputs" / "reports" / "test_classification_report.txt"
SCREENSHOT = Path(r"C:\Users\Lenovo\Pictures\Screenshots\屏幕截图 2026-06-22 151906.png")



def set_east_asia_font(obj, font_name):
    element = obj._element
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "F2F4F7")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Calibri"
    set_east_asia_font(r, "宋体")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)


def add_picture_if_exists(doc, path, caption, width=5.9):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)


def parse_last_log():
    rows = list(csv.DictReader(LOG_PATH.open("r", encoding="utf-8-sig")))
    last = rows[-1]
    best = max(rows, key=lambda r: float(r["val_acc"]))
    return rows, last, best


def extract_summary(report_path):
    text = report_path.read_text(encoding="utf-8")
    summary = {}
    for line in text.splitlines():
        parts = line.split()
        if not parts:
            continue
        if parts[0] == "accuracy":
            summary["accuracy"] = parts[1]
            summary["support"] = parts[2]
        elif parts[0] == "macro" and len(parts) >= 6:
            summary["macro_precision"] = parts[2]
            summary["macro_recall"] = parts[3]
            summary["macro_f1"] = parts[4]
        elif parts[0] == "weighted" and len(parts) >= 6:
            summary["weighted_precision"] = parts[2]
            summary["weighted_recall"] = parts[3]
            summary["weighted_f1"] = parts[4]
    return summary, text


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    set_east_asia_font(normal, "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        set_east_asia_font(style, "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《机器学习》结课论文")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Calibri"
    set_east_asia_font(r, "黑体")

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 ResNet18-FPN 的交通标志识别系统")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Calibri"
    set_east_asia_font(r, "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("——多尺度特征融合与 PyTorch 实现")
    r.font.size = Pt(14)
    r.font.name = "Calibri"
    set_east_asia_font(r, "黑体")

    for _ in range(7):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年6月22日")
    r.font.size = Pt(12)
    set_east_asia_font(r, "宋体")

    doc.add_page_break()


def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.29)
    return p


def main():
    rows, last, best = parse_last_log()
    val_summary, val_report_text = extract_summary(VAL_REPORT)
    test_summary, test_report_text = extract_summary(TEST_REPORT)

    doc = Document()
    setup_styles(doc)
    add_title_page(doc)

    doc.add_heading("摘 要", level=1)
    add_paragraph(doc, "交通标志识别是智能交通系统和辅助驾驶场景中的重要感知任务。旧版项目采用HOG特征结合SVM与随机森林完成传统机器学习流程，本项目在此基础上进一步设计并实现基于深度卷积神经网络的交通标志识别系统。项目使用GTSRB数据集，构建Small-Image ResNet18主干，并在其后加入FPN式多尺度特征融合结构，以缓解小尺寸交通标志在深层网络下采样过程中细节丢失的问题。")
    add_paragraph(doc, f"实验基于PyTorch实现，训练集包含39209张图像，按8:2进行分层划分得到7842张验证图像，并使用官方Test.csv对应的12630张测试图像进行独立评估。训练30轮后，模型最终训练准确率达到{float(last['train_acc'])*100:.2f}%，验证准确率达到{float(last['val_acc'])*100:.2f}%，最佳验证准确率为{float(best['val_acc'])*100:.2f}%。在官方测试集上，模型准确率为{float(test_summary['accuracy'])*100:.2f}%，加权F1分数为{float(test_summary['weighted_f1'])*100:.2f}%。实验结果表明，残差结构与多尺度特征融合能够有效提升交通标志分类任务中的特征表达能力。")

    doc.add_heading("1 业务问题定义", level=1)
    add_paragraph(doc, "交通标志识别是自动驾驶与高级驾驶辅助系统中环境感知的核心任务之一。交通标志承载着限速、禁止、警示、指示等道路规则信息，如果驾驶员因疲劳、遮挡或复杂光照条件漏看标志，可能带来安全风险。利用摄像头采集图像并自动识别交通标志类别，可以为车辆决策、驾驶提醒和道路场景理解提供基础能力。")
    add_paragraph(doc, "本项目关注的是已裁剪交通标志图像的多类别分类问题，即给定单张交通标志图片，模型需要判断其属于43个GTSRB类别中的哪一类。与完整道路场景中的检测任务不同，本项目不负责从街景中定位标志区域，而是聚焦于图像分类模型的设计、训练、评估和单张图片预测流程。")

    doc.add_heading("2 数据加载与探索性分析", level=1)
    add_paragraph(doc, "项目使用德国交通标志识别基准数据集GTSRB。该数据集训练部分包含39209张图像，覆盖43个类别；官方测试集包含12630张图像。数据集通过CSV文件记录图像路径和类别标签，Train.csv与Test.csv均包含Width、Height、ROI坐标、ClassId和Path等字段。")
    add_table(doc, ["数据部分", "样本数量", "用途"], [
        ["Train.csv", "39209", "训练与验证划分"],
        ["验证集", "7842", "从训练集按8:2分层划分得到"],
        ["Test.csv", "12630", "官方独立测试集评估"],
        ["类别数", "43", "GTSRB交通标志类别"],
    ], widths=[1.4, 1.2, 3.7])
    add_paragraph(doc, "交通标志图片原始尺寸并不统一，部分样本尺寸较小。为了保证批量训练的张量维度一致，同时避免小图像在网络前端过早损失空间细节，本项目将所有图像统一缩放到64×64×3后输入模型。")

    doc.add_heading("3 数据预处理与数据增强", level=1)
    add_paragraph(doc, "与旧版HOG+SVM项目不同，新版CNN模型不再手工提取HOG特征，而是让卷积层自动学习边缘、颜色、形状和语义特征。输入图像首先使用OpenCV读取并转换为RGB格式，再统一resize到64×64，随后将像素值归一化到0到1之间，并调整为PyTorch所需的C×H×W张量格式。")
    add_paragraph(doc, "训练阶段加入了轻量数据增强，包括小角度旋转、平移、缩放、亮度对比度扰动以及少量高斯噪声。由于交通标志包含左右方向、箭头等方向敏感类别，本项目没有使用水平翻转，以免引入错误标签。")

    doc.add_heading("4 模型结构与训练设计", level=1)
    add_paragraph(doc, "模型采用Small-Image ResNet18-FPN结构。普通ResNet最初面向224×224图像，开头使用7×7大卷积和最大池化。考虑到本项目图片尺寸较小，模型前端改为3×3、stride=1的Stem卷积，并取消早期MaxPool，从而保留更多边缘、颜色和数字细节。")
    add_table(doc, ["模块", "输出尺度", "说明"], [
        ["Stem", "64×64×64", "3×3卷积，stride=1，不使用早期MaxPool"],
        ["Stage1 / C1", "64×64×64", "2个BasicBlock，不下采样，保留浅层细节"],
        ["Stage2 / C2", "32×32×128", "2个BasicBlock，首块stride=2"],
        ["Stage3 / C3", "16×16×256", "2个BasicBlock，提取中层形状特征"],
        ["Stage4 / C4", "8×8×512", "2个BasicBlock，提取深层语义特征"],
        ["FPN融合", "F1/F2/F3", "C4逐级上采样并与C3、C2、C1拼接"],
        ["分类头", "43维logits", "多尺度GAP拼接后接全连接分类器"],
    ], widths=[1.25, 1.35, 3.7])
    add_paragraph(doc, "FPN融合部分借鉴YOLO和特征金字塔思想，将深层C4特征逐级上采样，并分别与浅层C3、C2、C1进行拼接。这样既保留了深层语义信息，也重新引入浅层高分辨率细节，适合小尺寸交通标志分类任务。")
    add_table(doc, ["训练参数", "取值"], [
        ["框架", "PyTorch"],
        ["输入尺寸", "64×64×3"],
        ["优化器", "AdamW"],
        ["初始学习率", "0.001"],
        ["权重衰减", "1e-4"],
        ["训练轮数", "30"],
        ["类别权重", "使用balanced class weight缓解类别不均衡"],
        ["设备", "CUDA GPU"],
    ], widths=[1.8, 4.5])

    doc.add_heading("5 模型训练与评估结果", level=1)
    add_paragraph(doc, "训练过程中，模型在第1轮的训练准确率为13.20%，验证准确率为26.28%；第3轮验证准确率已提升至95.996%，说明残差网络能够快速学习交通标志的结构化视觉特征。随着训练继续，模型在后期趋于稳定，学习率由ReduceLROnPlateau策略逐步降低。")
    add_table(doc, ["指标", "最终第30轮", "最佳记录"], [
        ["训练loss", f"{float(last['train_loss']):.4f}", "-"],
        ["训练准确率", f"{float(last['train_acc'])*100:.2f}%", "-"],
        ["验证loss", f"{float(last['val_loss']):.4f}", "-"],
        ["验证准确率", f"{float(last['val_acc'])*100:.2f}%", f"{float(best['val_acc'])*100:.2f}%"],
        ["当前学习率", last["lr"], "-"],
    ], widths=[1.8, 2.0, 2.0])
    add_picture_if_exists(doc, FIG_DIR / "training_curves.png", "图1 训练与验证集准确率、损失曲线", width=5.8)
    add_picture_if_exists(doc, SCREENSHOT, "图2 训练完成后的控制台输出结果", width=5.8)

    doc.add_heading("6 结果分析与可视化", level=1)
    add_paragraph(doc, "在内部验证集上，模型准确率达到99.97%，宏平均F1为99.98%，加权F1为99.97%。这说明模型不仅整体准确率高，而且在多数类别上具有均衡表现。少数误差主要集中在形状和图案高度相似的类别，例如不同限速标志或细小符号差异较弱的类别。")
    add_table(doc, ["评估集", "样本数", "Accuracy", "Macro F1", "Weighted F1"], [
        ["验证集", val_summary.get("support", "7842"), f"{float(val_summary['accuracy'])*100:.2f}%", f"{float(val_summary['macro_f1'])*100:.2f}%", f"{float(val_summary['weighted_f1'])*100:.2f}%"],
        ["官方测试集", test_summary.get("support", "12630"), f"{float(test_summary['accuracy'])*100:.2f}%", f"{float(test_summary['macro_f1'])*100:.2f}%", f"{float(test_summary['weighted_f1'])*100:.2f}%"],
    ], widths=[1.4, 1.1, 1.2, 1.2, 1.2])
    add_picture_if_exists(doc, FIG_DIR / "validation_confusion_matrix.png", "图3 验证集混淆矩阵", width=5.8)
    add_picture_if_exists(doc, FIG_DIR / "test_confusion_matrix.png", "图4 官方测试集混淆矩阵", width=5.8)
    add_paragraph(doc, "在官方测试集上，模型准确率为98.90%，加权F1为98.88%。相较旧版HOG+SVM在官方测试集上的91.08%准确率，新版CNN模型在独立测试数据上具有更强泛化能力。其原因在于CNN能够自动学习颜色、边缘、纹理和局部图案的组合特征，而FPN融合结构进一步改善了小尺寸细节在深层网络中的保留问题。")

    doc.add_heading("7 总结与展望", level=1)
    add_paragraph(doc, "本项目在旧版交通标志识别大作业的基础上，将传统HOG+SVM流程升级为基于深度学习的PyTorch实现。通过Small-Image ResNet18主干、残差连接和FPN多尺度特征融合，模型能够在GTSRB数据集上取得较高识别精度。训练结果显示，模型在验证集上达到99.97%的准确率，在官方测试集上达到98.90%的准确率，说明该结构对交通标志分类任务具有良好的判别能力和泛化能力。")
    add_paragraph(doc, "后续工作可以从三个方向继续改进：第一，加入更完整的单张图片预测可视化与误分类样本分析；第二，将分类模型扩展到检测任务，使其能够从真实道路图像中先定位标志再分类；第三，尝试轻量化模型或模型压缩技术，以便在嵌入式设备或车载平台上实时运行。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()