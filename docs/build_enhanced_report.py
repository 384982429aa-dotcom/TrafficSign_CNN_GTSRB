# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import math

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"E:\class_work\TrafficSign_CNN_GTSRB")
TEMPLATE = Path(r"C:\Users\Lenovo\Desktop\xeizuo\机器学习\机器学习结课报告模版.docx")
OUT = ROOT / "docs" / "机器学习结课报告_基于ResNet18-FPN的交通标志识别_完整版.docx"
FIG_DIR = ROOT / "outputs" / "figures"
DOC_FIG_DIR = ROOT / "docs" / "report_assets"
LOG_PATH = ROOT / "outputs" / "logs" / "training_log.csv"
VAL_REPORT = ROOT / "outputs" / "reports" / "validation_classification_report.txt"
TEST_REPORT = ROOT / "outputs" / "reports" / "test_classification_report.txt"
TRAIN_CSV = Path(r"E:\ML\GTSRB\data\Train.csv")
TEST_CSV = Path(r"E:\ML\GTSRB\data\Test.csv")
DATA_ROOT = Path(r"E:\ML\GTSRB")
SCREENSHOT = Path(r"C:\Users\Lenovo\Pictures\Screenshots\屏幕截图 2026-06-22 151906.png")

DOC_FIG_DIR.mkdir(parents=True, exist_ok=True)


def set_east_asia_font(run, font_name):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def clear_document_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    if "Title" in doc.styles:
        st = doc.styles["Title"]
        st.font.name = "Calibri"
        st.font.size = Pt(18)
        st.font.bold = True
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    for style_name in ["一级标题", "Heading 1"]:
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = "Calibri"
            st.font.size = Pt(15)
            st.font.bold = True
            st.font.color.rgb = RGBColor(0, 0, 0)
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            st.paragraph_format.space_before = Pt(12)
            st.paragraph_format.space_after = Pt(6)


def add_center_text(doc, text, size=12, bold=False, font="宋体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = "Calibri"
    set_east_asia_font(r, font)
    return p


def add_body_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.29)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    set_east_asia_font(r, "宋体")
    return p


def add_no_indent_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    set_east_asia_font(r, "宋体")
    return p


def add_heading1(doc, text):
    style = "一级标题" if "一级标题" in doc.styles else "Heading 1"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.bold = True
    set_east_asia_font(r, "黑体")
    return p


def set_table_borders(table, top=True, middle=True, bottom=True):
    # Clear all borders, then apply a classic three-line table style.
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                tag = qn(f"w:{edge}")
                element = borders.find(tag)
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)
                element.set(qn("w:val"), "nil")

    def apply_row_border(row, edge, size="12"):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:color"), "000000")

    if top:
        apply_row_border(table.rows[0], "top", "12")
    if middle:
        apply_row_border(table.rows[0], "bottom", "8")
    if bottom:
        apply_row_border(table.rows[-1], "bottom", "12")


def cell_text(cell, text, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.bold = bold
    set_east_asia_font(r, "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_three_line_table(doc, headers, rows, widths=None, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.font.size = Pt(10)
        r.bold = True
        set_east_asia_font(r, "宋体")
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True, center=True)
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, center=(len(str(value)) < 18))
            if widths:
                cells[i].width = Inches(widths[i])
    set_table_borders(table)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)
    set_east_asia_font(r, "宋体")


def add_picture(doc, path, caption, width=5.8):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)


def parse_log():
    rows = list(csv.DictReader(LOG_PATH.open("r", encoding="utf-8-sig")))
    last = rows[-1]
    best = max(rows, key=lambda r: float(r["val_acc"]))
    return rows, last, best


def report_summary(path):
    text = path.read_text(encoding="utf-8")
    s = {}
    class_rows = []
    for line in text.splitlines():
        parts = line.split()
        if not parts:
            continue
        if parts[0].isdigit() and len(parts) == 5:
            class_rows.append({
                "class": int(parts[0]),
                "precision": float(parts[1]),
                "recall": float(parts[2]),
                "f1": float(parts[3]),
                "support": int(parts[4]),
            })
        elif parts[0] == "accuracy":
            s["accuracy"] = parts[1]
            s["support"] = parts[2]
        elif parts[0] == "macro" and len(parts) >= 5:
            s["macro_f1"] = parts[4]
        elif parts[0] == "weighted" and len(parts) >= 5:
            s["weighted_f1"] = parts[4]
    return s, class_rows


def data_stats():
    train = pd.read_csv(TRAIN_CSV)
    test = pd.read_csv(TEST_CSV)
    counts = train["ClassId"].value_counts().sort_index()
    sizes = train[["Width", "Height"]]
    return {
        "train": train,
        "test": test,
        "counts": counts,
        "min_count": int(counts.min()),
        "max_count": int(counts.max()),
        "imbalance": float(counts.max() / counts.min()),
        "min_w": int(sizes["Width"].min()),
        "max_w": int(sizes["Width"].max()),
        "mean_w": float(sizes["Width"].mean()),
        "min_h": int(sizes["Height"].min()),
        "max_h": int(sizes["Height"].max()),
        "mean_h": float(sizes["Height"].mean()),
    }


def get_font(size=22, bold=False):
    candidates = [
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for font_path in candidates:
        if Path(font_path).exists():
            try:
                return ImageFont.truetype(font_path, size=size)
            except Exception:
                pass
    return ImageFont.load_default()


def draw_centered_text(draw, box, text, font, fill=(0, 0, 0)):
    x, y, w, h = box
    lines = text.split("\n")
    line_heights = []
    line_widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1] + 6)
    total_h = sum(line_heights)
    cy = y + (h - total_h) / 2
    for line, lw, lh in zip(lines, line_widths, line_heights):
        draw.text((x + (w - lw) / 2, cy), line, font=font, fill=fill)
        cy += lh


def make_architecture_diagram(path):
    img = Image.new("RGB", (1850, 1080), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(36)
    box_font = get_font(22)
    small_font = get_font(18)
    tiny_font = get_font(15)
    blue = (31, 77, 120)
    dark = (11, 37, 69)
    gray = (95, 95, 95)
    green = (45, 120, 70)
    orange = (130, 90, 30)
    fill_encoder = (232, 238, 245)
    fill_decoder = (235, 247, 238)
    fill_head = (250, 242, 222)
    fill_bridge = (240, 236, 250)
    fill_res = (245, 248, 252)

    draw.text((420, 35), "Small-Image ResNet18-FPN U-Shaped Multi-scale Fusion", font=title_font, fill=(0, 0, 0))
    draw.text((160, 105), "Encoder: ResNet18 backbone, downsample for semantic features", font=small_font, fill=blue)
    draw.text((1120, 105), "Decoder: FPN upsample + concatenate shallow details", font=small_font, fill=green)

    def box(x, y, w, h, text, fill, font=None):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=14, fill=fill, outline=blue, width=3)
        draw_centered_text(draw, (x, y, w, h), text, font or box_font)
        return (x, y, w, h)

    def c_left(b):
        x, y, w, h = b
        return (x, y + h / 2)

    def c_right(b):
        x, y, w, h = b
        return (x + w, y + h / 2)

    def c_top(b):
        x, y, w, h = b
        return (x + w / 2, y)

    def c_bottom(b):
        x, y, w, h = b
        return (x + w / 2, y + h)

    def arrow_head(p1, p2, color=dark, width=4, size=18):
        import math
        x1, y1 = p1
        x2, y2 = p2
        ang = math.atan2(y2 - y1, x2 - x1)
        for da in (2.55, -2.55):
            x = x2 - size * math.cos(ang + da)
            y = y2 - size * math.sin(ang + da)
            draw.line((x2, y2, x, y), fill=color, width=width)

    def arrow(p1, p2, color=dark, width=4):
        draw.line((p1, p2), fill=color, width=width)
        arrow_head(p1, p2, color=color, width=width)

    def elbow(points, color=dark, width=4):
        for a, b in zip(points, points[1:]):
            draw.line((a, b), fill=color, width=width)
        arrow_head(points[-2], points[-1], color=color, width=width)

    def skip(p1, p2, label):
        draw.line((p1, p2), fill=gray, width=3)
        x2, y2 = p2
        draw.ellipse((x2 - 9, y2 - 9, x2 + 9, y2 + 9), fill="white", outline=blue, width=3)
        draw.text(((p1[0] + p2[0]) / 2 - 35, p1[1] - 30), label, font=tiny_font, fill=gray)

    # Encoder: left side of the U.
    enc_x, enc_w, h = 95, 260, 100
    input_b = box(enc_x, 155, enc_w, h, "Input\n64x64x3", fill_encoder)
    stem_b = box(enc_x, 280, enc_w, h, "Stem\n3x3 Conv\n64x64x64", fill_encoder)
    c1_b = box(enc_x, 405, enc_w, h, "C1\nResBlock x2\n64x64x64", fill_encoder)
    c2_b = box(enc_x, 530, enc_w, h, "C2\nResBlock x2\n32x32x128", fill_encoder)
    c3_b = box(enc_x, 655, enc_w, h, "C3\nResBlock x2\n16x16x256", fill_encoder)
    c4_b = box(enc_x, 780, enc_w, h, "C4 / Bottleneck\nResBlock x2\n8x8x512", fill_bridge)

    for a, b in [(input_b, stem_b), (stem_b, c1_b), (c1_b, c2_b), (c2_b, c3_b), (c3_b, c4_b)]:
        arrow(c_bottom(a), c_top(b))

    # Decoder: right side of the U. FPN features go upward.
    dec_x, dec_w = 1120, 300
    f3_b = box(dec_x, 655, dec_w, h, "F3\nUpsample(C4) + C3\n16x16x256", fill_decoder)
    f2_b = box(dec_x, 530, dec_w, h, "F2\nUpsample(F3) + C2\n32x32x128", fill_decoder)
    f1_b = box(dec_x, 405, dec_w, h, "F1\nUpsample(F2) + C1\n64x64x64", fill_decoder)

    # C4 is upsampled directly below F3, then connected upward without a bend.
    up_b = box(1175, 799, 190, 62, "Upsample\n8x8 -> 16x16", fill_head, tiny_font)
    arrow(c_right(c4_b), c_left(up_b), color=dark, width=4)
    arrow(c_top(up_b), c_bottom(f3_b), color=dark, width=4)
    arrow(c_top(f3_b), c_bottom(f2_b))
    arrow(c_top(f2_b), c_bottom(f1_b))

    # Horizontal skip/concat links.
    skip(c_right(c3_b), c_left(f3_b), "Concat C3")
    skip(c_right(c2_b), c_left(f2_b), "Concat C2")
    skip(c_right(c1_b), c_left(f1_b), "Concat C1")

    # Classification head is outside the U, so the structure does not look tangled.
    gap_b = box(dec_x, 210, dec_w, 105, "GAP(F1,F2,F3,C4)\nMulti-scale concat", fill_head)
    fc_b = box(1480, 210, 150, 105, "Dense\nDropout", fill_head)
    out_b = box(1670, 210, 145, 105, "Softmax\n43 classes", fill_head)
    arrow(c_top(f1_b), c_bottom(gap_b))
    arrow(c_right(gap_b), c_left(fc_b))
    arrow(c_right(fc_b), c_left(out_b))
    draw.text((1180, 175), "Classification head", font=small_font, fill=orange)

    # Resolution labels.
    for y, label in [(445, "64x64"), (570, "32x32"), (695, "16x16"), (820, "8x8")]:
        draw.text((35, y), label, font=tiny_font, fill=(80, 80, 80))

    # ResBlock detail: right-bottom inset, drawn vertically.
    inset = (1460, 650, 355, 380)
    ix, iy, iw, ih = inset
    draw.rounded_rectangle((ix, iy, ix + iw, iy + ih), radius=16, fill=fill_res, outline=(160, 170, 185), width=2)
    draw.text((ix + 74, iy + 18), "Basic ResBlock", font=small_font, fill=(0, 0, 0))
    rb_font = get_font(15)
    block_x, block_w, block_h = ix + 118, 125, 54
    a = box(block_x, iy + 70, block_w, block_h, "3x3 Conv\nBN + ReLU", (255, 255, 255), rb_font)
    b = box(block_x, iy + 145, block_w, block_h, "3x3 Conv\nBN", (255, 255, 255), rb_font)
    add_node = (ix + 180, iy + 240)
    relu = box(block_x + 20, iy + 285, 85, 50, "ReLU", (255, 255, 255), rb_font)
    arrow(c_bottom(a), c_top(b), color=dark, width=3)
    arrow(c_bottom(b), (add_node[0], add_node[1] - 18), color=dark, width=3)
    draw.ellipse((add_node[0] - 18, add_node[1] - 18, add_node[0] + 18, add_node[1] + 18), fill=(255, 255, 255), outline=blue, width=3)
    draw.line((add_node[0] - 10, add_node[1], add_node[0] + 10, add_node[1]), fill=blue, width=3)
    draw.line((add_node[0], add_node[1] - 10, add_node[0], add_node[1] + 10), fill=blue, width=3)
    arrow((add_node[0], add_node[1] + 18), c_top(relu), color=dark, width=3)
    # Identity shortcut runs down the left side and joins at Add.
    shortcut_x = ix + 72
    shortcut_start = (block_x, iy + 97)
    elbow([shortcut_start, (shortcut_x, shortcut_start[1]), (shortcut_x, add_node[1]), (add_node[0] - 18, add_node[1])], color=(90, 90, 90), width=3)
    draw.text((ix + 62, iy + 345), "identity shortcut", font=tiny_font, fill=(80, 80, 80))
    # Legend.
    draw.rounded_rectangle((575, 900, 1180, 1005), radius=10, fill=(248, 248, 248), outline=(190, 190, 190), width=2)
    draw.rectangle((610, 928, 645, 951), fill=fill_encoder, outline=blue)
    draw.text((665, 921), "Encoder feature: downsample", font=small_font, fill=(0, 0, 0))
    draw.rectangle((610, 965, 645, 988), fill=fill_decoder, outline=blue)
    draw.text((665, 958), "Decoder feature: upsample + concat", font=small_font, fill=(0, 0, 0))
    draw.line((1010, 940, 1090, 940), fill=gray, width=3)
    draw.ellipse((1081, 931, 1099, 949), fill="white", outline=blue, width=3)
    draw.text((1110, 928), "skip / concat", font=small_font, fill=(0, 0, 0))

    img.save(path)

def make_data_distribution(stats, path):
    counts = stats["counts"]
    img = Image.new("RGB", (1500, 650), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(32)
    font = get_font(18)
    small = get_font(14)
    draw.text((430, 30), "GTSRB训练集类别分布", font=title_font, fill=(0,0,0))
    left, top, right, bottom = 90, 110, 1440, 540
    draw.line((left, bottom, right, bottom), fill=(0,0,0), width=2)
    draw.line((left, top, left, bottom), fill=(0,0,0), width=2)
    max_count = counts.max()
    bar_w = (right-left) / len(counts)
    for idx, (cls, count) in enumerate(counts.items()):
        x0 = left + idx * bar_w + 2
        x1 = left + (idx + 1) * bar_w - 2
        h = (count / max_count) * (bottom - top)
        y0 = bottom - h
        draw.rectangle((x0, y0, x1, bottom), fill=(46,116,181))
        if idx % 2 == 0:
            draw.text((x0, bottom + 8), str(cls), font=small, fill=(0,0,0))
    for ratio in [0, .25, .5, .75, 1.0]:
        y = bottom - ratio * (bottom-top)
        draw.line((left-5, y, right, y), fill=(220,220,220), width=1)
        draw.text((20, y-10), str(int(max_count*ratio)), font=small, fill=(0,0,0))
    draw.text((650, 590), "ClassId", font=font, fill=(0,0,0))
    draw.text((10, 90), "样本数", font=font, fill=(0,0,0))
    img.save(path)

def add_cover(doc):
    add_center_text(doc, "《机器学习》结课论文", 18, True, "黑体")
    for _ in range(5):
        doc.add_paragraph()
    add_center_text(doc, "基于 ResNet18-FPN 的交通标志识别系统", 20, True, "黑体")
    add_center_text(doc, "——多尺度特征融合与 PyTorch 实现", 14, False, "黑体")
    for _ in range(7):
        doc.add_paragraph()
    add_center_text(doc, "2026年6月22日", 12, False, "宋体")
    doc.add_page_break()


def add_abstract(doc, last, best, test):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("摘 要")
    r.bold = True
    r.font.size = Pt(14)
    set_east_asia_font(r, "黑体")
    add_body_para(doc, "交通标志识别是智能交通系统和辅助驾驶中的重要图像分类任务。本项目基于GTSRB交通标志数据集，设计并实现了一个基于PyTorch的卷积神经网络分类系统。针对交通标志图像尺寸较小、类别间形状相似、样本分布不均衡等问题，模型采用适合小图像输入的ResNet18残差主干，并结合FPN式多尺度特征融合结构，将浅层高分辨率细节与深层语义特征进行拼接融合。")
    add_body_para(doc, f"实验使用39209张训练图像，并按8:2划分训练集与验证集，同时使用官方Test.csv对应的12630张图像进行独立测试。训练30轮后，模型训练准确率为{float(last['train_acc'])*100:.2f}%，验证准确率为{float(last['val_acc'])*100:.2f}%，最佳验证准确率为{float(best['val_acc'])*100:.2f}%；在官方测试集上准确率为{float(test['accuracy'])*100:.2f}%，加权F1分数为{float(test['weighted_f1'])*100:.2f}%。结果表明，ResNet18-FPN结构能够有效学习交通标志的局部细节与整体语义特征，具有较好的分类精度和泛化能力。")
    add_no_indent_para(doc, "关键词：交通标志识别；GTSRB；ResNet18；FPN；多尺度特征融合；PyTorch")
    doc.add_paragraph()


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目录")
    r.bold = True
    r.font.size = Pt(14)
    set_east_asia_font(r, "黑体")
    for item in [
        "一、研究背景",
        "二、数据加载与探索性分析",
        "三、数据预处理与数据增强",
        "四、模型选型与训练",
        "五、模型调优与评估",
        "六、结果分析与可视化",
        "七、问题分析与解决方案",
        "八、总结与展望",
    ]:
        add_no_indent_para(doc, item)
    doc.add_page_break()


def main():
    rows, last, best = parse_log()
    val, val_classes = report_summary(VAL_REPORT)
    test, test_classes = report_summary(TEST_REPORT)
    stats = data_stats()
    arch_path = DOC_FIG_DIR / "resnet18_fpn_architecture.png"
    dist_path = DOC_FIG_DIR / "class_distribution.png"
    make_architecture_diagram(arch_path)
    make_data_distribution(stats, dist_path)

    worst_test = sorted(test_classes, key=lambda x: x["f1"])[:5]
    best_epoch = best["epoch"]

    doc = Document(str(TEMPLATE))
    clear_document_body(doc)
    setup_styles(doc)
    add_cover(doc)

    p = doc.add_paragraph(style="Title" if "Title" in doc.styles else None)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 ResNet18-FPN 的交通标志识别系统")
    r.bold = True
    r.font.size = Pt(18)
    set_east_asia_font(r, "黑体")

    add_abstract(doc, last, best, test)
    add_toc(doc)

    add_heading1(doc, "一、研究背景")
    add_body_para(doc, "交通标志是道路交通规则的重要载体，包含限速、禁止、警告和指示等信息。对于智能交通系统和辅助驾驶系统而言，自动识别交通标志能够帮助车辆理解道路环境，并为驾驶决策提供辅助依据。随着深度学习在计算机视觉领域的发展，卷积神经网络已经成为图像分类任务中的主流方法。")
    add_body_para(doc, "本项目面向已裁剪交通标志图像的分类任务，目标是输入一张交通标志图片，输出其所属类别。该任务的难点主要包括三点：第一，交通标志原始图像尺寸不统一，小尺寸图像较多；第二，不同类别之间颜色和外轮廓相似，区别可能只体现在中心数字、箭头或局部符号上；第三，各类别样本数量不均衡，少数类别样本较少，模型容易受到多数类影响。")
    add_body_para(doc, "因此，本项目选择设计一个适合小尺寸图像的残差卷积网络，并引入多尺度特征融合，使模型既能利用深层语义特征，也能保留浅层高分辨率细节。整个系统包含数据读取、数据分析、预处理、模型训练、评估可视化和单张图片预测等完整流程。")

    add_heading1(doc, "二、数据加载与探索性分析")
    add_body_para(doc, "项目使用德国交通标志识别基准数据集GTSRB。数据集根目录为E:\\ML\\GTSRB，其中data文件夹下包含Train.csv、Test.csv和Meta.csv。CSV文件中Path字段记录图像相对路径，ClassId字段记录交通标志类别编号。程序通过读取CSV并拼接数据集根目录得到完整图片路径。")
    add_three_line_table(doc, ["数据部分", "样本数量", "作用"], [
        ["Train.csv", "39209", "用于训练与验证划分"],
        ["验证集", "7842", "从训练集按8:2分层划分"],
        ["Test.csv", "12630", "官方独立测试集评估"],
        ["类别数", "43", "GTSRB全部交通标志类别"],
    ], widths=[1.4, 1.1, 3.8], caption="表1 数据集组成")
    add_body_para(doc, f"从类别分布看，训练集中样本最多的类别有{stats['max_count']}张，样本最少的类别仅有{stats['min_count']}张，最大类别与最小类别数量比约为{stats['imbalance']:.1f}:1，说明数据集存在明显类别不均衡。如果不做处理，模型可能更偏向样本数量较多的类别。")
    add_picture(doc, dist_path, "图1 GTSRB训练集类别分布", width=5.9)
    add_body_para(doc, "图1展示了43个类别在训练集中的样本数量差异。可以看到，不同类别之间样本数量差异明显，部分限速类和常见指示类样本较多，而一些少见警告标志样本较少。针对这一问题，训练阶段引入balanced class weight，使少数类在损失函数中获得更高权重。")
    add_body_para(doc, f"从图像尺寸看，训练图像宽度范围为{stats['min_w']}到{stats['max_w']}像素，平均宽度约为{stats['mean_w']:.1f}像素；高度范围为{stats['min_h']}到{stats['max_h']}像素，平均高度约为{stats['mean_h']:.1f}像素。原始图像尺寸较小且不统一，因此需要在输入模型前统一缩放。")

    add_heading1(doc, "三、数据预处理与数据增强")
    add_body_para(doc, "数据预处理首先使用OpenCV读取图片，并将BGR格式转换为RGB格式；随后将图像统一缩放为64×64，并将像素值除以255归一化到0到1之间，最后转换为PyTorch模型所需的C×H×W张量格式。统一尺寸能够保证批量训练时张量维度一致，而64×64相比48×48保留了更多局部细节，也不会带来过高计算成本。")
    add_body_para(doc, "训练阶段加入轻量级数据增强，包括小角度旋转、平移、缩放、亮度对比度扰动以及少量高斯噪声。这些增强操作模拟真实道路场景中常见的拍摄角度、光照变化和成像噪声，有助于提升模型泛化能力。需要注意的是，交通标志中存在方向敏感类别，例如左转、右转、靠左和靠右，因此本项目没有使用水平翻转，以免改变图像语义。")
    add_three_line_table(doc, ["处理步骤", "实现方式", "作用"], [
        ["读取图像", "cv2.imread + BGR转RGB", "统一颜色通道顺序"],
        ["尺寸统一", "resize到64×64", "适配批量训练并保留小图细节"],
        ["归一化", "像素值除以255", "稳定梯度与训练过程"],
        ["数据增强", "旋转、平移、缩放、亮度扰动、噪声", "提升鲁棒性"],
        ["标签映射", "ClassId映射为连续label_index", "适配交叉熵损失"],
    ], widths=[1.2, 2.0, 3.1], caption="表2 数据预处理流程")

    add_heading1(doc, "四、模型选型与训练")
    add_body_para(doc, "本项目采用Small-Image ResNet18-FPN模型。ResNet18通过残差连接缓解深层网络训练中的梯度衰减问题；FPN结构则将深层语义特征逐级上采样，并与浅层高分辨率特征进行拼接，使模型同时利用局部细节和全局语义。与原版ResNet不同，本项目面向64×64小图像，因此前端使用3×3、stride=1的Stem卷积，并取消早期大步长卷积和最大池化，避免小图像在网络前几层被过度压缩。")
    add_picture(doc, arch_path, "图2 Small-Image ResNet18-FPN模型结构图", width=6.2)
    add_body_para(doc, "图2给出了模型整体结构。输入图像经过Stem卷积后进入四个残差阶段，分别得到C1、C2、C3和C4四个尺度的特征图。其中C1分辨率最高，保留较多边缘和颜色细节；C4分辨率最低，但语义信息最强。FPN部分将C4逐级上采样，并分别与C3、C2、C1拼接，得到F3、F2和F1。最终对F1、F2、F3和C4进行全局平均池化并拼接为多尺度特征向量，输入全连接分类头输出43类logits。")
    add_three_line_table(doc, ["模块", "输出尺度", "说明"], [
        ["Stem", "64×64×64", "3×3卷积，stride=1，保留小图细节"],
        ["C1", "64×64×64", "浅层边缘和颜色特征"],
        ["C2", "32×32×128", "中低层形状特征"],
        ["C3", "16×16×256", "中高层图案特征"],
        ["C4", "8×8×512", "深层语义特征"],
        ["FPN", "F1/F2/F3", "上采样并拼接C1、C2、C3特征"],
        ["分类头", "43类", "多尺度全局池化后进行分类"],
    ], widths=[1.2, 1.3, 3.9], caption="表3 模型结构分解")
    add_three_line_table(doc, ["训练设置", "取值"], [
        ["深度学习框架", "PyTorch"],
        ["输入尺寸", "64×64×3"],
        ["优化器", "AdamW"],
        ["初始学习率", "0.001"],
        ["权重衰减", "1e-4"],
        ["训练轮数", "30"],
        ["类别权重", "balanced class weight"],
        ["运行设备", "CUDA GPU"],
    ], widths=[1.8, 4.4], caption="表4 训练超参数设置")

    add_heading1(doc, "五、模型调优与评估")
    add_body_para(doc, "训练过程中使用AdamW优化器并设置权重衰减，以减少过拟合风险。学习率调度器根据验证集准确率变化进行自适应调整，当验证性能长期不提升时降低学习率。损失函数采用交叉熵损失，并结合类别权重缓解类别样本数量不均衡问题。")
    add_three_line_table(doc, ["指标", "第30轮结果", "最佳结果"], [
        ["训练loss", f"{float(last['train_loss']):.4f}", "-"],
        ["训练准确率", f"{float(last['train_acc'])*100:.2f}%", "-"],
        ["验证loss", f"{float(last['val_loss']):.4f}", "-"],
        ["验证准确率", f"{float(last['val_acc'])*100:.2f}%", f"{float(best['val_acc'])*100:.2f}%（第{best_epoch}轮）"],
        ["学习率", last["lr"], "-"],
    ], widths=[1.8, 2.0, 2.0], caption="表5 训练过程核心指标")
    add_three_line_table(doc, ["控制台字段", "数值", "含义"], [
        ["Epoch", "030/30", "完成第30轮训练，总训练轮数为30"],
        ["train_loss", f"{float(last['train_loss']):.4f}", "训练集平均交叉熵损失"],
        ["train_acc", f"{float(last['train_acc'])*100:.2f}%", "训练集分类准确率"],
        ["val_loss", f"{float(last['val_loss']):.4f}", "验证集平均交叉熵损失"],
        ["val_acc", f"{float(last['val_acc'])*100:.2f}%", "验证集分类准确率"],
        ["Best validation accuracy", f"{float(best['val_acc'])*100:.2f}%", "训练过程中保存最优模型的依据"],
        ["Best model", "resnet18_fpn_best.pth", "验证集准确率最高的模型权重"],
        ["Final model", "resnet18_fpn_final.pth", "第30轮结束后的最终模型权重"],
    ], widths=[1.8, 1.8, 2.7], caption="表6 训练完成控制台输出解析")
    add_body_para(doc, "表6将训练结束时的控制台输出整理为表格。Best model表示根据验证集准确率保存的最优模型，Final model表示最后一轮训练完成后的模型。实际部署或单张图片预测时，一般优先加载验证集表现最好的resnet18_fpn_best.pth。")
    add_picture(doc, FIG_DIR / "training_curves.png", "图3 模型训练与验证曲线", width=5.9)
    add_body_para(doc, "图3左侧为训练准确率和验证准确率变化曲线，右侧为训练损失和验证损失变化曲线。模型在前3轮快速收敛，验证准确率迅速提升到95%以上，说明网络能够较快捕捉交通标志的主要视觉模式。后续曲线趋于平稳，训练损失和验证损失均接近0，表明模型已充分收敛。")
    add_picture(doc, SCREENSHOT, "图4 训练完成后的控制台输出截图", width=5.8)
    add_body_para(doc, "图4展示了训练完成后的终端输出，可以看到第30轮训练准确率达到99.94%，验证准确率达到99.97%，最佳验证准确率达到99.99%。这说明模型在训练集和验证集上都获得了很高的分类性能。")

    add_heading1(doc, "六、结果分析与可视化")
    add_body_para(doc, "为了全面评价模型性能，本项目分别在验证集和官方测试集上计算precision、recall、F1-score与accuracy。验证集来自训练集的分层划分，与训练数据同源；官方测试集独立提供，更能反映模型泛化性能。")
    add_three_line_table(doc, ["评估集", "样本数", "Accuracy", "Macro F1", "Weighted F1"], [
        ["验证集", val.get("support", "7842"), f"{float(val['accuracy'])*100:.2f}%", f"{float(val['macro_f1'])*100:.2f}%", f"{float(val['weighted_f1'])*100:.2f}%"],
        ["官方测试集", test.get("support", "12630"), f"{float(test['accuracy'])*100:.2f}%", f"{float(test['macro_f1'])*100:.2f}%", f"{float(test['weighted_f1'])*100:.2f}%"],
    ], widths=[1.3, 1.1, 1.2, 1.2, 1.2], caption="表7 验证集与官方测试集评估结果")
    add_body_para(doc, "表7显示，模型在验证集上的Accuracy为99.97%，在官方测试集上的Accuracy为98.90%。官方测试集性能略低于验证集，原因在于测试图像的光照、裁剪和拍摄条件与训练集存在一定差异。尽管如此，官方测试集的加权F1仍达到98.88%，说明模型具有较强的泛化能力。")
    add_picture(doc, FIG_DIR / "validation_confusion_matrix.png", "图5 验证集混淆矩阵", width=5.9)
    add_body_para(doc, "图5为验证集混淆矩阵。矩阵主对角线颜色最深，说明绝大多数样本被正确分类。验证集上仅有极少数样本发生误判，主要集中在外观非常接近的交通标志类别之间，例如不同限速数值或相似警示图案。")
    add_picture(doc, FIG_DIR / "test_confusion_matrix.png", "图6 官方测试集混淆矩阵", width=5.9)
    add_body_para(doc, "图6为官方测试集混淆矩阵。与验证集相比，测试集非对角线误判略有增加，说明独立测试数据中的光照和裁剪差异带来了一定挑战。但整体对角线仍然非常明显，证明模型在43类交通标志上具有稳定识别能力。")
    add_three_line_table(doc, ["ClassId", "Precision", "Recall", "F1-score", "Support"], [
        [r["class"], f"{r['precision']:.4f}", f"{r['recall']:.4f}", f"{r['f1']:.4f}", r["support"]] for r in worst_test
    ], widths=[1.0, 1.2, 1.2, 1.2, 1.0], caption="表8 官方测试集中F1较低的类别")
    add_body_para(doc, "表8列出了官方测试集中F1相对较低的若干类别。这些类别通常具有样本数量较少、图像细节较细或与其他类别外观相似等特点。例如部分警告标志只在内部符号上存在细微差别，当图像分辨率较低或存在模糊时，模型更容易出现混淆。")

    add_heading1(doc, "七、问题分析与解决方案")
    add_body_para(doc, "问题一是小尺寸图像在深层网络中容易因连续下采样而丢失细节。解决方案是在模型结构上采用小图像版本ResNet18，取消原始ResNet前端的大步长卷积和最大池化，并通过FPN多尺度拼接将浅层细节重新融合到分类头中。")
    add_body_para(doc, "问题二是类别不均衡。训练集中不同类别样本数量差异明显，如果直接训练，模型可能偏向样本较多类别。解决方案是在训练脚本中使用compute_class_weight计算类别权重，并传入CrossEntropyLoss，使少数类样本在损失中具有更高权重。")
    add_body_para(doc, "问题三是训练过程观察不直观。由于数据集规模较大，当batch size较小时一个epoch包含数千个batch，如果只在epoch结束时打印日志，会让人误以为程序卡住。因此训练脚本增加了--log-interval参数，支持按batch间隔输出当前loss和accuracy，便于实时观察训练状态。")
    add_body_para(doc, "问题四是外部图片预测路径不固定。为此项目新增了src\\prediction\\predict_image.py脚本，运行后在控制台提示用户输入图片地址，程序自动加载最优模型并输出预测类别、类别名称、置信度和Top-5结果，同时显示输入图片。")

    add_heading1(doc, "八、总结与展望")
    add_body_para(doc, "本项目完成了基于ResNet18-FPN的交通标志识别系统设计与实现。系统覆盖数据读取、探索性分析、图像预处理、数据增强、模型构建、训练、评估、模型保存以及单张图片预测等完整流程。实验结果表明，残差网络结合多尺度特征融合能够在GTSRB交通标志分类任务上取得较高准确率。")
    add_body_para(doc, "从结果看，模型在验证集上达到99.97%的准确率，在官方测试集上达到98.90%的准确率，说明模型不仅能够充分拟合训练分布，也能较好适应独立测试集。FPN结构对小尺寸交通标志细节保留具有积极作用，类别权重和数据增强也提升了模型鲁棒性。")
    add_body_para(doc, "后续可以从三个方向继续改进：第一，将当前分类模型扩展为检测加分类的完整交通标志识别系统，使其能够处理真实道路场景图像；第二，尝试轻量化网络、模型剪枝或量化，以便部署到边缘设备；第三，收集更多复杂光照、遮挡和运动模糊条件下的样本，进一步评估和提升模型在真实环境中的泛化能力。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

