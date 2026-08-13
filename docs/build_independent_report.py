# -*- coding: utf-8 -*-
from pathlib import Path
import csv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"E:\class_work\TrafficSign_CNN_GTSRB")
TEMPLATE = Path(r"C:\Users\Lenovo\Desktop\xeizuo\机器学习\机器学习结课报告模版.docx")
OUT = ROOT / "docs" / "机器学习结课报告_基于ResNet18-FPN的交通标志识别.docx"
FIG_DIR = ROOT / "outputs" / "figures"
LOG_PATH = ROOT / "outputs" / "logs" / "training_log.csv"
VAL_REPORT = ROOT / "outputs" / "reports" / "validation_classification_report.txt"
TEST_REPORT = ROOT / "outputs" / "reports" / "test_classification_report.txt"
SCREENSHOT = Path(r"C:\Users\Lenovo\Pictures\Screenshots\屏幕截图 2026-06-22 151906.png")


def set_east_asia_font(run, font_name):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def clear_document_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for style_name in ["Title", "标题"]:
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = "Calibri"
            st.font.size = Pt(18)
            st.font.bold = True
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    for style_name in ["一级标题", "Heading 1"]:
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = "Calibri"
            st.font.size = Pt(15)
            st.font.bold = True
            st.font.color.rgb = RGBColor(0, 0, 0)
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            st.paragraph_format.space_before = Pt(12)
            st.paragraph_format.space_after = Pt(6)


def add_center_text(doc, text, size=12, bold=False, font="宋体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = "Calibri"
    set_east_asia_font(r, font)
    return p


def add_body_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.29)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    set_east_asia_font(r, "宋体")
    return p


def add_heading1(doc, text):
    style = "一级标题" if "一级标题" in doc.styles else "Heading 1"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.bold = True
    set_east_asia_font(r, "黑体")
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.bold = bold
    set_east_asia_font(r, "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = False
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)
    set_east_asia_font(r, "宋体")


def add_picture(doc, path, caption, width=5.8):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)


def parse_log():
    rows = list(csv.DictReader(LOG_PATH.open("r", encoding="utf-8-sig")))
    last = rows[-1]
    best = max(rows, key=lambda r: float(r["val_acc"]))
    return rows, last, best


def report_summary(path):
    text = path.read_text(encoding="utf-8")
    s = {}
    for line in text.splitlines():
        parts = line.split()
        if not parts:
            continue
        if parts[0] == "accuracy":
            s["accuracy"] = parts[1]
            s["support"] = parts[2]
        elif parts[0] == "macro" and len(parts) >= 5:
            s["macro_f1"] = parts[4]
        elif parts[0] == "weighted" and len(parts) >= 5:
            s["weighted_f1"] = parts[4]
    return s


def add_cover(doc):
    add_center_text(doc, "《机器学习》结课论文", 18, True, "黑体")
    for _ in range(5):
        doc.add_paragraph()
    add_center_text(doc, "基于 ResNet18-FPN 的交通标志识别系统", 20, True, "黑体")
    add_center_text(doc, "——多尺度特征融合与 PyTorch 实现", 14, False, "黑体")
    for _ in range(7):
        doc.add_paragraph()
    add_center_text(doc, "2026年6月22日", 12, False, "宋体")
    doc.add_page_break()


def add_abstract(doc, last, best, test):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("摘 要")
    r.bold = True
    r.font.size = Pt(14)
    set_east_asia_font(r, "黑体")
    add_body_para(doc, "交通标志识别是智能交通系统和辅助驾驶中的重要图像分类任务。本项目基于GTSRB交通标志数据集，设计并实现了一个基于PyTorch的卷积神经网络分类系统。针对交通标志图像尺寸较小、类别间形状相似、样本分布不均衡等问题，模型采用适合小图像输入的ResNet18残差主干，并结合FPN式多尺度特征融合结构，将浅层高分辨率细节与深层语义特征进行拼接融合。")
    add_body_para(doc, f"实验使用39209张训练图像，并按8:2划分训练集与验证集，同时使用官方Test.csv对应的12630张图像进行独立测试。训练30轮后，模型训练准确率为{float(last['train_acc'])*100:.2f}%，验证准确率为{float(last['val_acc'])*100:.2f}%，最佳验证准确率为{float(best['val_acc'])*100:.2f}%；在官方测试集上准确率为{float(test['accuracy'])*100:.2f}%，加权F1分数为{float(test['weighted_f1'])*100:.2f}%。结果表明，ResNet18-FPN结构能够有效学习交通标志的局部细节与整体语义特征，具有较好的分类精度和泛化能力。")
    doc.add_paragraph()


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目录")
    r.bold = True
    r.font.size = Pt(14)
    set_east_asia_font(r, "黑体")
    items = [
        "一、研究背景",
        "二、数据加载与探索性分析",
        "三、数据预处理与数据增强",
        "四、模型选型与训练",
        "五、模型调优与评估",
        "六、结果分析与可视化",
        "七、问题分析与解决方案",
        "八、总结与展望",
    ]
    for item in items:
        add_body_para(doc, item)
    doc.add_page_break()


def main():
    rows, last, best = parse_log()
    val = report_summary(VAL_REPORT)
    test = report_summary(TEST_REPORT)

    doc = Document(str(TEMPLATE))
    clear_document_body(doc)
    setup_styles(doc)

    add_cover(doc)
    p = doc.add_paragraph(style="Title" if "Title" in doc.styles else None)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 ResNet18-FPN 的交通标志识别系统")
    r.bold = True
    r.font.size = Pt(18)
    set_east_asia_font(r, "黑体")
    add_abstract(doc, last, best, test)
    add_toc(doc)

    add_heading1(doc, "一、研究背景")
    add_body_para(doc, "交通标志是道路交通规则的重要载体，包含限速、禁止、警告和指示等信息。对于智能交通系统和辅助驾驶系统而言，自动识别交通标志能够帮助车辆理解道路环境，并为驾驶决策提供辅助依据。随着深度学习在计算机视觉领域的发展，卷积神经网络已成为图像分类任务中的主流方法。")
    add_body_para(doc, "本项目面向已裁剪交通标志图像的分类任务，目标是输入一张交通标志图片，输出其所属类别。该任务的难点在于图像尺寸不统一、部分类别样本数量较少、不同类别之间存在高度相似的颜色和几何结构，同时拍摄角度、光照和清晰度也会影响识别效果。")

    add_heading1(doc, "二、数据加载与探索性分析")
    add_body_para(doc, "项目使用德国交通标志识别基准数据集GTSRB。数据集根目录为E:\\ML\\GTSRB，其中data文件夹下包含Train.csv、Test.csv和Meta.csv。CSV文件中Path字段记录图像相对路径，ClassId字段记录交通标志类别编号。")
    add_table(doc, ["数据部分", "样本数量", "作用"], [
        ["Train.csv", "39209", "用于训练与验证划分"],
        ["验证集", "7842", "从训练集按8:2分层划分"],
        ["Test.csv", "12630", "官方独立测试集评估"],
        ["类别数", "43", "GTSRB全部交通标志类别"],
    ], widths=[1.4, 1.1, 3.8])
    add_body_para(doc, "从数据结构看，训练图像按照类别文件夹组织，测试图像集中存放。图像原始尺寸从二十多像素到一百多像素不等，因此需要在输入模型前统一尺寸。")

    add_heading1(doc, "三、数据预处理与数据增强")
    add_body_para(doc, "数据预处理首先使用OpenCV读取图像，并将BGR格式转换为RGB格式；随后将图像统一缩放为64×64，并将像素值除以255归一化到0到1之间，最后转换为PyTorch模型所需的C×H×W张量。")
    add_body_para(doc, "训练阶段加入随机旋转、平移、缩放、亮度对比度扰动和轻微噪声，以提升模型在不同拍摄条件下的鲁棒性。由于交通标志中存在方向敏感类别，例如左转、右转等，数据增强中没有采用水平翻转，避免产生语义错误样本。")

    add_heading1(doc, "四、模型选型与训练")
    add_body_para(doc, "本项目采用Small-Image ResNet18-FPN模型。ResNet18通过残差连接缓解深层网络训练中的梯度衰减问题；FPN结构则将深层语义特征逐级上采样，并与浅层高分辨率特征进行拼接，使模型同时利用局部细节和全局语义。")
    add_table(doc, ["模块", "输出尺度", "说明"], [
        ["Stem", "64×64×64", "3×3卷积，stride=1，保留小图细节"],
        ["C1", "64×64×64", "浅层边缘和颜色特征"],
        ["C2", "32×32×128", "中低层形状特征"],
        ["C3", "16×16×256", "中高层图案特征"],
        ["C4", "8×8×512", "深层语义特征"],
        ["FPN", "F1/F2/F3", "上采样并拼接C1、C2、C3特征"],
        ["分类头", "43类", "多尺度全局池化后Softmax分类"],
    ], widths=[1.2, 1.3, 3.9])
    add_table(doc, ["训练设置", "取值"], [
        ["深度学习框架", "PyTorch"],
        ["输入尺寸", "64×64×3"],
        ["优化器", "AdamW"],
        ["初始学习率", "0.001"],
        ["训练轮数", "30"],
        ["类别权重", "balanced class weight"],
        ["运行设备", "CUDA GPU"],
    ], widths=[1.8, 4.4])

    add_heading1(doc, "五、模型调优与评估")
    add_body_para(doc, "训练过程中使用AdamW优化器并设置权重衰减，以减少过拟合风险。学习率调度器根据验证集准确率变化进行自适应调整，当验证性能长期不提升时降低学习率。损失函数采用交叉熵损失，并结合类别权重缓解类别样本数量不均衡问题。")
    add_table(doc, ["指标", "第30轮结果", "最佳结果"], [
        ["训练loss", f"{float(last['train_loss']):.4f}", "-"],
        ["训练准确率", f"{float(last['train_acc'])*100:.2f}%", "-"],
        ["验证loss", f"{float(last['val_loss']):.4f}", "-"],
        ["验证准确率", f"{float(last['val_acc'])*100:.2f}%", f"{float(best['val_acc'])*100:.2f}%"],
        ["学习率", last["lr"], "-"],
    ], widths=[1.8, 2.0, 2.0])
    add_picture(doc, FIG_DIR / "training_curves.png", "图1 模型训练与验证曲线", width=5.8)
    add_picture(doc, SCREENSHOT, "图2 训练完成后的控制台输出", width=5.8)

    add_heading1(doc, "六、结果分析与可视化")
    add_body_para(doc, "从训练曲线可以看出，模型在前几轮快速收敛，第3轮验证准确率已达到较高水平，后续训练逐步逼近稳定状态。最终验证集准确率达到99.97%，官方测试集准确率达到98.90%，说明模型在同源验证集和独立测试集上均具有较好的识别能力。")
    add_table(doc, ["评估集", "样本数", "Accuracy", "Macro F1", "Weighted F1"], [
        ["验证集", val.get("support", "7842"), f"{float(val['accuracy'])*100:.2f}%", f"{float(val['macro_f1'])*100:.2f}%", f"{float(val['weighted_f1'])*100:.2f}%"],
        ["官方测试集", test.get("support", "12630"), f"{float(test['accuracy'])*100:.2f}%", f"{float(test['macro_f1'])*100:.2f}%", f"{float(test['weighted_f1'])*100:.2f}%"],
    ], widths=[1.3, 1.1, 1.2, 1.2, 1.2])
    add_picture(doc, FIG_DIR / "validation_confusion_matrix.png", "图3 验证集混淆矩阵", width=5.8)
    add_picture(doc, FIG_DIR / "test_confusion_matrix.png", "图4 官方测试集混淆矩阵", width=5.8)

    add_heading1(doc, "七、问题分析与解决方案")
    add_body_para(doc, "本项目训练初期容易出现控制台长时间无输出的现象，原因是脚本最初只在一个epoch结束后打印日志，而batch size较小时每个epoch包含大量batch。为了解决这一问题，训练脚本增加了--log-interval参数，使程序能够按batch间隔输出loss和accuracy，便于观察训练是否正常进行。")
    add_body_para(doc, "另一个问题是小尺寸图像在深层网络中容易因连续下采样而丢失细节。模型结构上采用小图像版本ResNet18，取消早期大步长卷积和最大池化，并引入FPN多尺度拼接，将浅层细节特征重新融合到分类头中，从结构上增强对小目标细节的保留。")
    add_body_para(doc, "此外，类别样本数量不均衡会导致模型偏向样本较多的类别。训练脚本通过compute_class_weight计算类别权重，并传入交叉熵损失函数，使少数类在损失中获得更高权重，从而提升整体类别均衡性。")

    add_heading1(doc, "八、总结与展望")
    add_body_para(doc, "本项目完成了基于ResNet18-FPN的交通标志识别系统设计与实现。系统覆盖数据读取、预处理、数据增强、模型构建、训练、评估、模型保存以及单张图片预测等完整流程。实验结果表明，残差网络结合多尺度特征融合能够在GTSRB交通标志分类任务上取得较高准确率。")
    add_body_para(doc, "后续可以继续从三个方向改进：第一，将分类模型扩展为检测加分类的完整交通标志识别系统；第二，进一步研究轻量化网络和模型压缩，以便部署到边缘设备；第三，增加更多真实场景图片测试，评估模型在复杂光照、遮挡和运动模糊条件下的鲁棒性。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()